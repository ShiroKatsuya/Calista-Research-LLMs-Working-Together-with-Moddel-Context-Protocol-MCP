import os
from langchain_core.tools import Tool
from langchain_google_community import GoogleSearchAPIWrapper
from langchain_core.documents import Document
from langchain_community.document_loaders import WebBaseLoader
import re
import textwrap
import urllib3
from bs4 import BeautifulSoup
import requests
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
import heapq
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from collections import Counter
import io
import tempfile
import os.path
from urllib.parse import urlparse, unquote


import fitz  
import docx2txt
import pdfplumber
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None


try:
    import pytesseract
    from PIL import Image
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False


try:
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('punkt')
    nltk.download('stopwords')

os.environ["GOOGLE_CSE_ID"] = "5650b4d81ac344bd8"
os.environ["GOOGLE_API_KEY"] = "AIzaSyBEl28l8XlkhYWxCkGamWLAWW2Jfc3SJR0"
os.environ["USER_AGENT"] = "jfE7aT4JW0dDcqCLyoiVfQ==kblDif2SgvOKHDhQ"
search = GoogleSearchAPIWrapper(k=5)
# results = tool.run("apa dampak jika tidur hanya 4 jam sehari")

def query_google(query):
    # Modified to search for academic content 
    academic_query = query
    # Check if we need to add academic terms
    if not any(term in query.lower() for term in ["pdf", "research", "journal", "thesis", "dissertation", "paper", "academia", "article"]):
        academic_query = f"{query} (pdf OR research paper OR journal OR dissertation OR thesis)"
    
    results = search.results(academic_query, 5)
    urls = []
    for result in results:
        urls.append(result['link'])
    print(urls)
    return urls

# Helper function to detect file type from content and URL
def detect_document_type(url, response=None):
    """
    Detect the type of document based on URL, content-type headers, and content
    
    Args:
        url (str): The URL of the document
        response (requests.Response, optional): The response object if already fetched
        
    Returns:
        str: Document type ('pdf', 'docx', 'doc', 'html', 'txt', etc.)
    """
    # Check file extension in URL
    parsed_url = urlparse(url)
    path = unquote(parsed_url.path.lower())
    
    if path.endswith('.pdf'):
        return 'pdf'
    elif path.endswith('.docx'):
        return 'docx'
    elif path.endswith('.doc'):
        return 'doc'
    elif path.endswith('.txt'):
        return 'txt'
    elif path.endswith('.rtf'):
        return 'rtf'
    elif any(path.endswith(ext) for ext in ['.ppt', '.pptx']):
        return 'presentation'
    elif any(path.endswith(ext) for ext in ['.xls', '.xlsx', '.csv']):
        return 'spreadsheet'
    
    # Check content-type header if response is provided
    if response and 'Content-Type' in response.headers:
        content_type = response.headers['Content-Type'].lower()
        if 'application/pdf' in content_type:
            return 'pdf'
        elif 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' in content_type:
            return 'docx'
        elif 'application/msword' in content_type:
            return 'doc'
        elif 'text/plain' in content_type:
            return 'txt'
        elif 'text/html' in content_type:
            return 'html'
        elif 'application/rtf' in content_type:
            return 'rtf'
    
    # Default to HTML if we can't determine the type
    return 'html'

# Extract text from PDF using PyMuPDF (fitz)
def extract_text_from_pdf(pdf_content):
    """
    Extract text content from a PDF using PyMuPDF
    
    Args:
        pdf_content (bytes): The binary content of the PDF
        
    Returns:
        str: Extracted text from the PDF
    """
    try:
        # Open PDF from binary content
        pdf_document = fitz.open(stream=pdf_content, filetype="pdf")
        
        # Extract text from each page
        text = []
        for page_num in range(pdf_document.page_count):
            page = pdf_document[page_num]
            page_text = page.get_text()
            text.append(f"\n--- Page {page_num + 1} ---\n{page_text}")
            
            # Check if page has low text content, might be scanned - try OCR if available
            if TESSERACT_AVAILABLE and len(page_text.strip()) < 100:
                try:
                    # Render page to image
                    pix = page.get_pixmap()
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    
                    # Use OCR to extract text
                    ocr_text = pytesseract.image_to_string(img)
                    if ocr_text.strip():
                        text.append(f"\n[OCR Text from Page {page_num + 1}]:\n{ocr_text}")
                except Exception as e:
                    print(f"OCR error on page {page_num + 1}: {e}")
        
        pdf_document.close()
        return "\n".join(text)
        
    except Exception as e:
        # Fallback to pdfplumber if fitz fails
        print(f"PyMuPDF extraction failed, trying pdfplumber: {e}")
        try:
            text = []
            with pdfplumber.open(io.BytesIO(pdf_content)) as pdf:
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""
                    text.append(f"\n--- Page {i + 1} ---\n{page_text}")
            return "\n".join(text)
        except Exception as e2:
            print(f"All PDF extraction methods failed: {e2}")
            return f"[Error extracting PDF content: {e2}]"

# Extract text from Word documents
def extract_text_from_docx(docx_content):
    """
    Extract text from a docx file
    
    Args:
        docx_content (bytes): The binary content of the docx file
        
    Returns:
        str: Extracted text from the docx file
    """
    try:
        # First try docx2txt which works well for most documents
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.write(docx_content)
            tmp_path = tmp.name
        
        text = docx2txt.process(tmp_path)
        os.unlink(tmp_path)  # Delete the temp file
        
        if text and len(text) > 100:
            return text
            
        # If docx2txt didn't work well, try python-docx
        if DocxDocument is not None:
            doc = DocxDocument(io.BytesIO(docx_content))
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # Handle tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
                        
            return '\n'.join(full_text)
        
        return text or "[No text content extracted]"
        
    except Exception as e:
        print(f"Error extracting text from docx: {e}")
        return f"[Error extracting Word document content: {e}]"

# Function to detect if URL is likely from an academic source
def is_academic_source(url):
    """
    Check if a URL is likely from an academic or scholarly source
    
    Args:
        url (str): The URL to check
        
    Returns:
        bool: True if likely academic, False otherwise
    """
    academic_domains = [
        '.edu', '.ac.', '.gov', 'scholar.google', 'researchgate.net', 
        'academia.edu', 'jstor.org', 'sciencedirect.com', 'springer.com',
        'ieee.org', 'ncbi.nlm.nih.gov', 'pubmed', 'arxiv.org',
        'ssrn.com', 'nature.com', 'science.org', 'wiley.com', 'tandfonline.com',
        'sagepub.com', 'elsevier.com', 'acm.org', 'apa.org'
    ]
    
    academic_keywords = [
        'journal', 'research', 'proceedings', 'conference', 'thesis',
        'dissertation', 'paper', 'study', 'article', 'peer-review'
    ]
    
    # Check domain
    if any(domain in url.lower() for domain in academic_domains):
        return True
        
    # Check for PDF links with academic keywords
    if url.lower().endswith('.pdf'):
        decoded_url = unquote(url.lower())
        if any(keyword in decoded_url for keyword in academic_keywords):
            return True
            
    return False

# Extract text from a document URL
def extract_text_from_document_url(url, query):
    """
    Extract text from a document at the specified URL
    
    Args:
        url (str): The URL of the document
        query (str): The search query for highlighting relevant content
        
    Returns:
        tuple: (extracted_text, document_type, is_academic)
    """
    try:
        # Headers to mimic a browser
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml,application/pdf,*/*',
            'Accept-Language': 'en-US,en;q=0.9'
        }
        
        # Download the content
        response = requests.get(url, headers=headers, verify=False, timeout=30, stream=True)
        response.raise_for_status()
        
        # Detect document type
        doc_type = detect_document_type(url, response)
        
        # Check if academic source
        is_academic = is_academic_source(url)
        
        # Extra metadata to include
        metadata = {
            "source": url,
            "document_type": doc_type,
            "is_academic": is_academic
        }
        
        # Process based on document type
        if doc_type == 'pdf':
            content = extract_text_from_pdf(response.content)
            return content, doc_type, is_academic
            
        elif doc_type in ['docx', 'doc']:
            content = extract_text_from_docx(response.content)
            return content, doc_type, is_academic
            
        elif doc_type == 'txt':
            return response.text, doc_type, is_academic
            
        elif doc_type == 'html':
            # Use existing HTML extraction method
            content = extract_main_content(response.text, url)
            return content, doc_type, is_academic
            
        else:
            # For unsupported formats, just return a message
            return f"[Document type '{doc_type}' is not supported for extraction]", doc_type, is_academic
            
    except Exception as e:
        print(f"Error extracting text from {url}: {e}")
        return f"[Error: {str(e)}]", "unknown", False

def extract_main_content(html_content, url):

    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Remove common non-content elements
    for element in soup.find_all(['header', 'footer', 'nav', 'aside', 'script', 'style', 'noscript', 
                                 'iframe', 'form', 'button', 'svg']):
        element.decompose()
        
    # Remove common ad, navigation, and footer classes/IDs
    for selector in [
        '[class*="nav"]', '[class*="menu"]', '[class*="header"]', '[class*="footer"]',
        '[class*="banner"]', '[class*="ad-"]', '[class*="ads-"]', '[class*="widget"]',
        '[class*="sidebar"]', '[class*="comment"]', '[id*="nav"]', '[id*="menu"]',
        '[id*="header"]', '[id*="footer"]', '[id*="banner"]', '[id*="ad-"]', 
        '[id*="sidebar"]', '[id*="comment"]'
    ]:
        for element in soup.select(selector):
            element.decompose()
    

    main_content = None
    content_elements = soup.select('article, [class*="content"], [class*="post"], [class*="article"], main, #content, .content, .post, .article')
    
    if content_elements:
        # Use the largest content block (likely the main article)
        main_content = max(content_elements, key=lambda x: len(x.get_text(strip=True)))
    else:
        # Fallback: find div with the most paragraphs
        divs = soup.find_all('div')
        if divs:
            main_content = max(divs, key=lambda x: len(x.find_all('p')))
        else:
            main_content = soup.body
    
    if not main_content:
        return "Could not extract main content"
    
    # Extract text from main content, preserving only necessary HTML elements
    text = ""
    for element in main_content.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'blockquote']):
        element_text = element.get_text(strip=True)
        if element_text:
            if element.name.startswith('h'):
                text += f"\n### {element_text} ###\n"
            elif element.name == 'li':
                text += f"  • {element_text}\n"
            elif element.name == 'blockquote':
                text += f"> {element_text}\n"
            else:
                text += f"{element_text}\n\n"
    
    return text.strip()

def score_paragraph_relevance(paragraphs, query):

    if not paragraphs:
        return []
    

    corpus = [query] + paragraphs
    

    vectorizer = TfidfVectorizer(stop_words='english')
    tfidf_matrix = vectorizer.fit_transform(corpus)
    

    query_vector = tfidf_matrix[0:1]
    paragraph_vectors = tfidf_matrix[1:]
    similarity_scores = cosine_similarity(query_vector, paragraph_vectors)[0]
    
    # Pair paragraphs with their scores and sort by score
    scored_paragraphs = list(zip(paragraphs, similarity_scores))
    scored_paragraphs.sort(key=lambda x: x[1], reverse=True)
    
    return scored_paragraphs

def extract_keywords_from_query(query, num_keywords=8):

    # Tokenize the query
    try:
        # Try to determine if query is in Indonesian
        indonesian_markers = ['mengapa', 'kenapa', 'bagaimana', 'apa', 'siapa', 'kapan', 'dimana', 
                             'di', 'dan', 'atau', 'dengan', 'pada', 'untuk', 'dari', 'dalam',
                             'yang', 'adalah', 'ini', 'itu', 'oleh', 'jika', 'maka']
        
        is_indonesian = any(marker in query.lower().split() for marker in indonesian_markers)
        
        # Get stopwords for filtering - use English by default
        stop_words = set(stopwords.words('english'))
        
        # Add some common words that aren't useful as keywords
        additional_stops = {'why', 'how', 'what', 'when', 'where', 'who', 'which', 
                           'is', 'are', 'am', 'was', 'were', 'be', 'been', 'being',
                           'have', 'has', 'had', 'do', 'does', 'did', 'can', 'could',
                           'will', 'would', 'shall', 'should', 'may', 'might', 'must',
                           'and', 'or', 'but', 'if', 'then', 'else', 'when', 'up', 'down',
                           'in', 'out', 'on', 'off', 'over', 'under', 'again', 'further',
                           'then', 'once', 'here', 'there', 'all', 'any', 'both', 'each',
                           'few', 'more', 'most', 'other', 'some', 'such'}
        
        # Add Indonesian stopwords if the query seems to be in Indonesian
        if is_indonesian:
            indonesian_stopwords = indonesian_markers + [
                'bahwa', 'sebagai', 'ia', 'mereka', 'saya', 'kamu', 'kita', 'kami',
                'dia', 'nya', 'itu', 'ini', 'ke', 'ya', 'juga', 'dapat', 'akan', 'masih',
                'telah', 'harus', 'secara', 'tetapi', 'tidak', 'belum', 'lebih', 'sangat',
                'bisa', 'tersebut', 'tentang', 'demikian', 'ketika', 'namun', 'menjadi',
                'seperti', 'sehingga', 'hingga'
            ]
            additional_stops.update(indonesian_stopwords)
            
        stop_words.update(additional_stops)
        
        # Tokenize and filter out stopwords
        words = word_tokenize(query.lower())
        keywords = [word for word in words if word.isalnum() and word not in stop_words and len(word) > 2]
        
        # Count word frequencies
        word_counts = Counter(keywords)
        
        # Get the most common words
        common_words = [word for word, count in word_counts.most_common(num_keywords)]
        
        # If we don't have enough keywords, try to extract n-grams (2-word phrases)
        if len(common_words) < 3:
            # Get bigrams
            tokens = [token.lower() for token in words if token.isalnum()]
            bigrams = [" ".join(tokens[i:i+2]) for i in range(len(tokens)-1)]
            bigram_counts = Counter(bigrams)
            top_bigrams = [bg for bg, count in bigram_counts.most_common(3) if len(bg) > 5]
            
            # Add individual words from bigrams
            for bigram in top_bigrams:
                words_in_bigram = bigram.split()
                for word in words_in_bigram:
                    if word not in common_words and word not in stop_words and len(word) > 2:
                        common_words.append(word)
        
        return common_words
    except Exception as e:
        print(f"Keyword extraction error: {e}")
        # Fallback to simple extraction if NLP processing fails
        words = re.findall(r'\b[a-zA-Z]{3,}\b', query.lower())
        return list(set(words))[:num_keywords]

def extract_keywords_from_text(text, num_keywords=15):
    """
    Extract important keywords from text using TF-IDF.
    
    Args:
        text (str): The text to extract keywords from
        num_keywords (int): Maximum number of keywords to extract
        
    Returns:
        list: List of extracted keywords
    """
    try:
        # Check if text appears to be in Indonesian
        indonesian_markers = ['dan', 'yang', 'di', 'pada', 'dengan', 'untuk', 'dari', 'dalam',
                             'adalah', 'ini', 'itu', 'oleh', 'jika', 'maka',
                             'tidak', 'bisa', 'akan', 'telah', 'sudah', 'harus']
        
        # Count occurrences of Indonesian markers
        indonesian_count = 0
        for marker in indonesian_markers:
            indonesian_count += text.lower().count(' ' + marker + ' ')
        
        # Use English stopwords by default
        stop_words = set(stopwords.words('english'))
        
        # Add Indonesian stopwords if enough markers are found
        if indonesian_count > 3:
            indonesian_stopwords = indonesian_markers + [
                'bahwa', 'sebagai', 'ia', 'mereka', 'saya', 'kamu', 'kita', 'kami',
                'dia', 'nya', 'itu', 'ini', 'ke', 'ya', 'juga', 'dapat', 'akan', 'masih',
                'telah', 'harus', 'secara', 'tetapi', 'tidak', 'belum', 'lebih', 'sangat',
                'bisa', 'tersebut', 'tentang', 'demikian', 'ketika', 'namun', 'menjadi',
                'seperti', 'sehingga', 'hingga'
            ]
            stop_words.update(indonesian_stopwords)
        
        # Create a TF-IDF vectorizer
        vectorizer = TfidfVectorizer(
            max_df=0.9,
            min_df=1,
            max_features=500,
            stop_words=stop_words
        )
        
        # Create a dummy document if text is too short
        if len(text.split()) < 20:
            # Just use word frequency for short texts
            words = word_tokenize(text.lower())
            keywords = [word for word in words if word.isalnum() and word not in stop_words and len(word) > 2]
            word_counts = Counter(keywords)
            return [word for word, count in word_counts.most_common(num_keywords)]
        
        # Split text into sentences
        sentences = sent_tokenize(text)
        
        # Make sure there are enough sentences
        if len(sentences) < 2:
            # Split by newlines or punctuation if not enough sentences
            sentences = re.split(r'[.!?;:\n]+', text)
            sentences = [s.strip() for s in sentences if len(s.strip()) > 10]
        
        # Fit the vectorizer
        try:
            tfidf_matrix = vectorizer.fit_transform(sentences)
            
            # Get feature names (terms)
            feature_names = vectorizer.get_feature_names_out()
            
            # Sum up TF-IDF scores for each term across all sentences
            tfidf_scores = np.sum(tfidf_matrix.toarray(), axis=0)
            
            # Create a dictionary of terms and their scores
            term_scores = {term: score for term, score in zip(feature_names, tfidf_scores)}
            
            # Sort terms by score and take the top ones
            sorted_terms = sorted(term_scores.items(), key=lambda x: x[1], reverse=True)
            
            # Extract only the terms (without scores)
            top_keywords = [term for term, score in sorted_terms[:num_keywords] if len(term) > 2]
            
            return top_keywords
        except ValueError:
            # Fallback if vectorizer fails
            words = word_tokenize(text.lower())
            words = [word for word in words if word.isalnum() and word not in stop_words and len(word) > 2]
            word_counts = Counter(words)
            return [word for word, count in word_counts.most_common(num_keywords)]
    except Exception as e:
        print(f"Text keyword extraction error: {e}")
        # Ultimate fallback
        words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
        return list(set(words))[:num_keywords]

def filter_relevant_content(content, query, threshold=0.15, max_paragraphs=10):


    paragraphs = re.split(r'\n\s*\n', content)
    paragraphs = [p.strip() for p in paragraphs if p.strip()]
    

    if len(paragraphs) < 5:
        return content
    

    scored_paragraphs = score_paragraph_relevance(paragraphs, query)
    

    query_keywords = extract_keywords_from_query(query)
    

    top_content = "\n".join([p for p, score in scored_paragraphs[:5] if score > 0.05])
    content_keywords = extract_keywords_from_text(top_content)

    combined_keywords = list(set(query_keywords + content_keywords))
    

    combined_keywords = [kw for kw in combined_keywords if len(kw) > 2]
    
    # Log keywords for debugging
    print(f"Automatically extracted keywords for filtering: {', '.join(combined_keywords)}")

    relevant_paragraphs = []
    for p, score in scored_paragraphs:

        keyword_count = sum(1 for keyword in combined_keywords if keyword.lower() in p.lower())

        if score >= threshold or keyword_count >= 2:
            relevant_paragraphs.append(p)
    

    relevant_paragraphs = relevant_paragraphs[:max_paragraphs]
    

    header_indices = []
    for i, paragraph in enumerate(paragraphs):
        if paragraph.startswith('###'):
            header_indices.append(i)
    

    for i in header_indices:

        next_header = next((h for h in header_indices if h > i), len(paragraphs))
        has_relevant_content = any(p in relevant_paragraphs for p in paragraphs[i+1:next_header])
        
        if has_relevant_content and paragraphs[i] not in relevant_paragraphs:
            relevant_paragraphs.append(paragraphs[i])
    
    # Re-sort paragraphs to maintain original order
    paragraph_indices = {p: i for i, p in enumerate(paragraphs)}
    relevant_paragraphs.sort(key=lambda p: paragraph_indices.get(p, 0))
    
    return "\n\n".join(relevant_paragraphs)

def generate_summary(text, query, num_sentences=5):

    # Skip summarization if text is too short
    if len(text.split()) < 100:
        return text
        
    # Tokenize the text into sentences
    sentences = sent_tokenize(text)
    

    if len(sentences) <= num_sentences:
        return text
    
    # Clean sentences
    clean_sentences = [re.sub(r'[^\w\s]', '', s.lower()) for s in sentences]
    
    # Create TF-IDF vectorizer and calculate sentence scores
    vectorizer = TfidfVectorizer(stop_words='english')
    
    try:
        # Add query to the corpus to calculate relevance
        corpus = clean_sentences + [query.lower()]
        tfidf_matrix = vectorizer.fit_transform(corpus)
        
        # Calculate similarity between query and each sentence
        query_vector = tfidf_matrix[-1]
        sentence_vectors = tfidf_matrix[:-1]
        similarities = cosine_similarity(query_vector, sentence_vectors)[0]
        
        # Get indices of top sentences based on similarity score
        top_indices = heapq.nlargest(num_sentences, 
                                     range(len(similarities)), 
                                     similarities.__getitem__)
        
        # Sort indices to maintain original order
        top_indices.sort()
        
        # Create summary
        summary = [sentences[i] for i in top_indices]
        
        return " ".join(summary)
    except:
        # Fallback to simple extraction if vectorization fails
        return " ".join(sentences[:num_sentences])

def search_and_load(query):
    # Get URLs from Google search
    urls = query_google(query)
    
    # Suppress only the specific SSL verification warnings
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
    
    documents = []
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'}
    
    for url in urls:
        try:
            # Extract text based on document type (PDF, Word, HTML, etc.)
            content, doc_type, is_academic = extract_text_from_document_url(url, query)
            
            # Skip empty or very short content
            if not content or len(content) < 50:
                print(f"Skipping {url}: insufficient content")
                continue
                
            # Get paragraph scores for highlighting
            paragraphs = re.split(r'\n\s*\n', content)
            paragraphs = [p.strip() for p in paragraphs if p.strip()]
            scored_paragraphs = score_paragraph_relevance(paragraphs, query)
            
            # Create a score map for later highlighting
            score_map = {p: score for p, score in scored_paragraphs}
            
            # Filter content for relevance to the search query
            # For academic papers/PDFs, be less aggressive with filtering to preserve structure
            if doc_type in ['pdf', 'docx', 'doc'] and is_academic:
                threshold = 0.10  # Lower threshold for academic content
                max_paragraphs = 20  # Keep more content from academic sources
            else:
                threshold = 0.15
                max_paragraphs = 10
                
            relevant_content = filter_relevant_content(content, query, threshold, max_paragraphs)
            
            if relevant_content:
                # Generate summary focusing on content related to query
                summary = generate_summary(relevant_content, query)
                
                documents.append(Document(
                    page_content=relevant_content, 
                    metadata={
                        "source": url,
                        "document_type": doc_type,
                        "is_academic": is_academic,
                        "score_map": score_map,
                        "summary": summary
                    })
                )
            
        except Exception as e:
            print(f"Error processing {url}: {e}")
    
    # Fallback to WebBaseLoader if no content was extracted
    if not documents:
        try:
            fallback_loader = WebBaseLoader(
                web_paths=urls,
                verify_ssl=False
            )
            fallback_documents = fallback_loader.load()
            filtered_documents = []
            
            # Filter fallback documents for relevance
            for doc in fallback_documents:
                if doc.page_content.strip():
                    # Get paragraph scores for highlighting
                    paragraphs = re.split(r'\n\s*\n', doc.page_content)
                    paragraphs = [p.strip() for p in paragraphs if p.strip()]
                    scored_paragraphs = score_paragraph_relevance(paragraphs, query)
                    
                    # Create a score map for later highlighting
                    score_map = {p: score for p, score in scored_paragraphs}
                    
                    relevant_content = filter_relevant_content(doc.page_content, query)
                    if relevant_content:
                        # Generate summary focusing on content related to query
                        summary = generate_summary(relevant_content, query)
                        
                        filtered_documents.append(Document(
                            page_content=relevant_content, 
                            metadata={
                                **doc.metadata,
                                "document_type": "html",
                                "is_academic": False,
                                "score_map": score_map,
                                "summary": summary
                            })
                        )
            
            if filtered_documents:
                documents = filtered_documents
            else:
                documents = [Document(page_content="No relevant content found", metadata={"source": urls[0] if urls else "No URL", "document_type": "unknown", "is_academic": False})]
        except Exception as fallback_error:
            print(f"Error loading fallback documents: {fallback_error}")
            documents = [Document(page_content=f"Error: {str(fallback_error)}", metadata={"source": urls[0] if urls else "No URL", "document_type": "unknown", "is_academic": False})]
    
    # Sort documents by their average paragraph relevance score
    for doc in documents:
        if "score_map" in doc.metadata and doc.metadata["score_map"]:
            doc.metadata["avg_score"] = sum(doc.metadata["score_map"].values()) / len(doc.metadata["score_map"])
        else:
            doc.metadata["avg_score"] = 0
    
    # Sort documents by relevance score, with a preference for academic sources
    documents.sort(key=lambda x: (x.metadata.get("avg_score", 0) * (1.2 if x.metadata.get("is_academic", False) else 1.0)), reverse=True)
    
    return documents

def format_content(content, score_map=None):

    content = content.replace('\r\n', '\n').strip()
    
    # Remove extra whitespace
    content = re.sub(r'\n{3,}', '\n\n', content)
    
    # Remove very short lines (likely menu items, breadcrumbs, etc.)
    content = re.sub(r'(?m)^.{1,20}$\n', '', content)
    
    # Remove any remaining header/footer-like text
    content = re.sub(r'(?i)^(home|menu|navigation|contact us|about us|privacy policy|terms of service)[\s\n]*$', '', content, flags=re.MULTILINE)
    
    # Enhanced formatting
    paragraphs = re.split(r'\n\s*\n', content)
    formatted_paragraphs = []
    
    horizontal_line = "─" * 80
    
    for para in paragraphs:
        if not para.strip():
            continue
        
        # Get relevance score for this paragraph if available
        relevance_score = score_map.get(para.strip(), 0) if score_map else 0
        high_relevance = relevance_score >= 0.25  # Threshold for high relevance highlighting
            
        # Format headers with proper styling
        if para.startswith('###'):
            header_text = para.strip('#').strip()
            formatted_paragraphs.append(f"\n{horizontal_line}")
            formatted_paragraphs.append(f"📌 {header_text.upper()} 📌")
            formatted_paragraphs.append(f"{horizontal_line}")
            continue
        
        # Format subheaders (detect potential subheadings)
        elif re.match(r'^[A-Z][A-Za-z\s]{5,50}:$', para.strip()) or para.strip().isupper():
            formatted_paragraphs.append(f"▶ {para.strip()} ◀")
            continue
            
        # Improve list formatting
        elif para.strip().startswith('•') or re.match(r'^\s*[\d\-\*]+\.?\s', para.strip()):
            list_items = re.split(r'\n', para.strip())
            for item in list_items:
                # Standardize list bullets and add indentation
                item = re.sub(r'^\s*[\d\-\*]+\.?\s', '• ', item.strip())
                if not item.startswith('•'):
                    item = f"• {item}"
                    
                # Highlight highly relevant list items
                if high_relevance:
                    formatted_paragraphs.append(f"  🔍 {item}")
                else:
                    formatted_paragraphs.append(f"  {item}")
            continue
            
        # Enhance quote formatting
        elif para.strip().startswith('>'):
            quote_text = para.strip('> ').strip()
            wrapped_quote = textwrap.fill(quote_text, width=76)
            quote_lines = wrapped_quote.split('\n')
            formatted_quote = ['┌' + '─' * 78 + '┐']
            for line in quote_lines:
                formatted_quote.append(f"│ \"{line}\"" + ' ' * (76 - len(line)) + ' │')
            formatted_quote.append('└' + '─' * 78 + '┘')
            formatted_paragraphs.append('\n'.join(formatted_quote))
            continue
            
        # Highlight important information (keywords)
        important_keywords = ['important', 'note', 'warning', 'caution', 'key', 'critical', 'essential']
        has_important = any(keyword in para.lower() for keyword in important_keywords)
        
        # Format normal paragraphs
        single_line = re.sub(r'\s+', ' ', para.strip())
        
        # Add emphasis based on importance and relevance
        if has_important:
            wrapped = textwrap.fill(single_line, width=76)
            formatted_paragraphs.append(f"❗ {wrapped} ❗")
        elif high_relevance:
            # Highlight highly relevant content with a magnifying glass
            wrapped = textwrap.fill(single_line, width=75)
            formatted_paragraphs.append(f"🔍 {wrapped}")
        else:
            wrapped = textwrap.fill(single_line, width=80)
            formatted_paragraphs.append(wrapped)
    
    # Add proper spacing
    formatted_text = '\n\n'.join(formatted_paragraphs)
    
    # Format URLs in text for better visibility
    formatted_text = re.sub(r'(https?://\S+)', r'🔗 \1', formatted_text)
    
    return formatted_text

# Create tool that combines search and web loading.
tool = Tool(
    name="google_search_and_load",
    description="Search Google and load webpage content from results",
    func=search_and_load
)

# New function to extract bibliographic information from academic content
def extract_bibliographic_info(content, doc_type, is_academic):
    """
    Extract bibliographic information from academic content if available
    
    Args:
        content (str): The document content
        doc_type (str): The document type (pdf, docx, etc)
        is_academic (bool): Whether the document is from an academic source
        
    Returns:
        dict: Dictionary with bibliographic information if found
    """
    if not is_academic:
        return None
    
    biblio_info = {
        "title": None,
        "authors": None,
        "publication": None,
        "year": None,
        "doi": None
    }
    
    # Try to extract title - often at the beginning or in a specific format
    title_patterns = [
        r'(?i)(?:title|paper):\s*([^\n\.]+)',
        r'^\s*#\s+([^\n\.]+)',  # Markdown title
        r'^\s*([A-Z][^\.]+?)\s*\n',  # First sentence in all caps or title case
        r'(?<=abstract)\s*([^\n\.]+)'  # Title before abstract
    ]
    
    for pattern in title_patterns:
        match = re.search(pattern, content[:500])
        if match:
            title = match.group(1).strip()
            if len(title) > 10 and len(title) < 200:  # Reasonable title length
                biblio_info["title"] = title
                break
    
    # Try to extract authors
    author_patterns = [
        r'(?i)(?:author|by)(?:s)?:\s*([^\n\.]+)',
        r'(?i)([A-Za-z\s\.,]+(?:et al\.)?)\s*\([0-9]{4}\)',  # Author(s) followed by year
        r'(?i)([A-Za-z\s\.,]+)(?:,|\sand\s)(?:[A-Za-z\s\.,]+)(?:\.|,)'  # Multiple authors separated by commas or 'and'
    ]
    
    for pattern in author_patterns:
        match = re.search(pattern, content[:1000])
        if match:
            authors = match.group(1).strip()
            if 3 < len(authors) < 200:  # Reasonable author string length
                biblio_info["authors"] = authors
                break
    
    # Try to extract year
    year_match = re.search(r'(?:19|20)[0-9]{2}', content[:1000])
    if year_match:
        biblio_info["year"] = year_match.group(0)
    
    # Try to extract DOI
    doi_match = re.search(r'(?i)(?:doi|https?://doi\.org/)[\s:]*(10\.[0-9]{4,}[^\s\n\.]+)', content)
    if doi_match:
        biblio_info["doi"] = doi_match.group(1)
    
    # Try to extract publication name
    pub_patterns = [
        r'(?i)(?:journal|conference|proceedings)[\s:]+([^\n\.]+)',
        r'(?i)in\s+([^,\.\n]+?),\s+(?:vol|volume|pp|pages)'
    ]
    
    for pattern in pub_patterns:
        match = re.search(pattern, content[:2000])
        if match:
            publication = match.group(1).strip()
            if 5 < len(publication) < 200:  # Reasonable publication name length
                biblio_info["publication"] = publication
                break
    
    # Only return if we found at least some information
    if any(value for value in biblio_info.values()):
        return biblio_info
    return None

# Update the print_formatted_results function to include bibliographic information
def print_formatted_results(results):
    """
    Prints the formatted results in a more appealing way.
    
    Args:
        results: List of Document objects with content to format and display
    """
    print("\n" + "=" * 100)
    print(f"{'SEARCH RESULTS':^100}")
    print("=" * 100 + "\n")
    
    for i, doc in enumerate(results, 1):
        source = doc.metadata.get('source', 'Unknown')
        avg_score = doc.metadata.get('avg_score', 0)
        score_map = doc.metadata.get('score_map', {})
        summary = doc.metadata.get('summary', '')
        doc_type = doc.metadata.get('document_type', 'html').upper()
        is_academic = doc.metadata.get('is_academic', False)
        
        # Document type icon
        type_icon = "📄"  # Default
        if doc_type == 'PDF':
            type_icon = "📑"
        elif doc_type in ('DOCX', 'DOC'):
            type_icon = "📝"
        elif doc_type == 'PRESENTATION':
            type_icon = "🎭"
        elif doc_type == 'SPREADSHEET':
            type_icon = "📊"
            
        # Academic badge
        academic_badge = "🎓 ACADEMIC SOURCE" if is_academic else ""
        
        print(f"\n{'▓' * 100}")
        print(f"{type_icon} DOCUMENT {i} | TYPE: {doc_type} | {academic_badge}")
        print(f"SOURCE: {source} | RELEVANCE: {avg_score:.2f}")
        
        # Extract and display bibliographic information if available for academic sources
        if is_academic:
            biblio_info = extract_bibliographic_info(doc.page_content, doc_type.lower(), is_academic)
            if biblio_info:
                print(f"{'▓' * 100}\n")
                print("📚 BIBLIOGRAPHIC INFORMATION:")
                if biblio_info.get("title"):
                    print(f"TITLE: {biblio_info['title']}")
                if biblio_info.get("authors"):
                    print(f"AUTHORS: {biblio_info['authors']}")
                if biblio_info.get("publication"):
                    print(f"PUBLICATION: {biblio_info['publication']}")
                if biblio_info.get("year"):
                    print(f"YEAR: {biblio_info['year']}")
                if biblio_info.get("doi"):
                    print(f"DOI: {biblio_info['doi']}")
                print()
            else:
                print(f"{'▓' * 100}\n")
        else:
            print(f"{'▓' * 100}\n")
        
        # Print summary in a highlighted box if available
        if summary and summary != doc.page_content:
            print("📋 KEY POINTS:")
            print("┌" + "─" * 98 + "┐")
            
            # Wrap summary text to fit in the box
            wrapped_summary = textwrap.fill(summary, width=96)
            for line in wrapped_summary.split('\n'):
                print(f"│ {line:<96} │")
                
            print("└" + "─" * 98 + "┘\n")
        
        # Format content based on document type
        if doc_type in ('PDF', 'DOCX', 'DOC') and is_academic:
            # Use special formatting for academic papers to preserve structure
            content = doc.page_content
            content = content.replace('\r\n', '\n').strip()
            
            # Format page markers in PDFs
            content = re.sub(r'--- Page (\d+) ---', r'📄 PAGE \1 ', content)
            
            # Preserve the structure more for academic content
            print(content)
        else:
            # Use regular formatting for web content
            print(format_content(doc.page_content, score_map))
            
        print("\n" + "▬" * 100 + "\n")

# Example usage.
results = tool.run("apa dampak jika tidur hanya 4 jam sehari")
print_formatted_results(results)
