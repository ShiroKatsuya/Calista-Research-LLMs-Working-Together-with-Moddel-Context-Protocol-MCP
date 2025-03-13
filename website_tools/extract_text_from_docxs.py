import io
import tempfile
import os
import docx2txt
from docx import Document as DocxDocument
# Extract text from Word documents
def extract_text_from_docx(docx_content):
    """
    Extract text from a docx file
    
    Args:
        docx_content (bytes): The binary content of the docx file
        
    Returns:
        str: Extracted text from the docx file
    """
    try:
        # First try docx2txt which works well for most documents
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.write(docx_content)
            tmp_path = tmp.name
        
        text = docx2txt.process(tmp_path)
        os.unlink(tmp_path)  # Delete the temp file
        
        if text and len(text) > 100:
            return text
            
        # If docx2txt didn't work well, try python-docx
        if DocxDocument is not None:
            doc = DocxDocument(io.BytesIO(docx_content))
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # Handle tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
                        
            return '\n'.join(full_text)
        
        return text or "[No text content extracted]"
        
    except Exception as e:
        print(f"Error extracting text from docx: {e}")
        return f"[Error extracting Word document content: {e}]"